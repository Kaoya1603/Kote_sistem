from docxtpl import DocxTemplate
from docx.shared import Inches
import sqlite3


def compose_file():
    con = sqlite3.connect('timu3.db')
    cur = con.cursor()

    yonghu_shuliang = len(cur.execute('''SELECT * FROM 'пользователи' ''').fetchall())
    huowu = len(cur.execute('''SELECT * FROM 'номенклатура' ''').fetchall())
    shiyongde_huowu = len(cur.execute('''SELECT * FROM 'номенклатура' WHERE используется = 'да';''').fetchall())
    shijian = list(map(lambda x: x[0], cur.execute('''SELECT "время аренды" FROM 'оплата аренды';''').fetchall()))
    zhongdeng_shijian = sum(shijian) // len(shijian)

    doc = DocxTemplate("tpl.docx")

    doc.add_paragraph('Подпись администратора')
    doc.add_picture('Факсимиле АДМИН.png', width=Inches(4.0), height=Inches(.7))

    doc.add_paragraph('Печать')
    doc.add_picture('1.png', width=Inches(4.0))

    context = {
        'yonghu_shuliang': yonghu_shuliang,
        'huowu': huowu,
        'shiyongde_huowu': shiyongde_huowu,
        'zhongdeng_shijian': zhongdeng_shijian
    }

    doc.render(context)
    doc.save("res.docx")
